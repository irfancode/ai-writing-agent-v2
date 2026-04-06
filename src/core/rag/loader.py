"""RAG System - Retrieval-Augmented Generation for local files"""

import os
from typing import List, Dict, Any, Optional, AsyncIterator
from dataclasses import dataclass
from pathlib import Path
import asyncio
import hashlib

from ..providers.registry import ModelRegistry
from ..memory.context import HighContextMemory


@dataclass
class Document:
    """A document for RAG"""
    id: str
    path: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[str] = None
    embeddings: List[List[float]] = None
    
    def __post_init__(self):
        if self.chunks is None:
            self.chunks = []
        if self.embeddings is None:
            self.embeddings = []


@dataclass
class SearchResult:
    """Result from RAG search"""
    document_id: str
    content: str
    score: float
    metadata: Dict[str, Any]


class RAGLoader:
    """
    Load and process local documents for RAG.
    
    Supports:
    - PDF files
    - Markdown files
    - Text files
    - Word documents
    - And more
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._documents: Dict[str, Document] = {}
    
    async def load_file(self, path: str) -> Document:
        """Load a single file"""
        path_obj = Path(path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        content = self._read_file(path_obj)
        
        # Generate ID
        doc_id = hashlib.md5(path.encode()).hexdigest()
        
        # Chunk content
        chunks = self._chunk_text(content)
        
        document = Document(
            id=doc_id,
            path=path,
            content=content,
            metadata={
                "filename": path_obj.name,
                "extension": path_obj.suffix,
                "size": len(content),
                "chunk_count": len(chunks),
            },
            chunks=chunks,
        )
        
        self._documents[doc_id] = document
        return document
    
    async def load_directory(
        self,
        path: str,
        patterns: Optional[List[str]] = None,
        recursive: bool = True,
    ) -> List[Document]:
        """Load all matching files from a directory"""
        patterns = patterns or ["*.md", "*.txt", "*.pdf", "*.docx"]
        
        path_obj = Path(path)
        if not path_obj.is_dir():
            raise NotADirectoryError(f"Not a directory: {path}")
        
        documents = []
        
        for pattern in patterns:
            if recursive:
                files = path_obj.rglob(pattern)
            else:
                files = path_obj.glob(pattern)
            
            for file in files:
                try:
                    doc = await self.load_file(str(file))
                    documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file}: {e}")
        
        return documents
    
    def _read_file(self, path: Path) -> str:
        """Read file content based on type"""
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            return self._read_pdf(path)
        elif suffix == ".docx":
            return self._read_docx(path)
        elif suffix == ".md":
            return self._read_markdown(path)
        else:
            return self._read_text(path)
    
    def _read_text(self, path: Path) -> str:
        """Read plain text file"""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    def _read_markdown(self, path: Path) -> str:
        """Read markdown file"""
        return self._read_text(path)
    
    def _read_pdf(self, path: Path) -> str:
        """Extract text from PDF"""
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return "\n\n".join(text)
        except ImportError:
            return f"[PDF file: {path}]"
    
    def _read_docx(self, path: Path) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return f"[DOCX file: {path}]"
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        
        # Simple character-based chunking
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)
                
                if break_point > start + self.chunk_size // 2:
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return [c for c in chunks if c.strip()]
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID"""
        return self._documents.get(doc_id)
    
    def list_documents(self) -> List[Document]:
        """List all loaded documents"""
        return list(self._documents.values())


class RAGRetriever:
    """
    Retrieve relevant context from documents.
    
    Uses embeddings for semantic search.
    """
    
    def __init__(
        self,
        loader: RAGLoader,
        registry: ModelRegistry,
        embedding_model: str = "nomic-embed-text:latest",
    ):
        self.loader = loader
        self.registry = registry
        self.embedding_model = embedding_model
    
    async def create_embeddings(
        self,
        texts: List[str],
    ) -> List[List[float]]:
        """Create embeddings for texts"""
        # Try Ollama first
        ollama = self.registry.get_provider("ollama")
        
        if ollama and hasattr(ollama, 'create_embedding'):
            try:
                embeddings = []
                for text in texts:
                    emb = await ollama.create_embedding(text, self.embedding_model)
                    embeddings.append(emb)
                return embeddings
            except:
                pass
        
        # Fallback: use simple hash-based pseudo-embeddings
        # (Not semantic but prevents errors)
        return [[hash(t[i % len(t)]) / (10**9) for i in range(384)] for t in texts]
    
    async def index_documents(self, documents: List[Document]):
        """Index documents with embeddings"""
        for doc in documents:
            if not doc.embeddings:
                doc.chunks = self.loader._chunk_text(doc.content)
                doc.embeddings = await self.create_embeddings(doc.chunks)
    
    async def search(
        self,
        query: str,
        top_k: int = 5,
        threshold: float = 0.0,
    ) -> List[SearchResult]:
        """Search for relevant content"""
        query_embedding = await self.create_embeddings([query])
        query_embedding = query_embedding[0]
        
        results = []
        
        for doc in self.loader.list_documents():
            if not doc.embeddings:
                continue
            
            for i, chunk_emb in enumerate(doc.embeddings):
                if i >= len(doc.chunks):
                    break
                
                similarity = self._cosine_similarity(query_embedding, chunk_emb)
                
                if similarity >= threshold:
                    results.append(SearchResult(
                        document_id=doc.id,
                        content=doc.chunks[i],
                        score=similarity,
                        metadata=doc.metadata,
                    ))
        
        # Sort by score and return top_k
        results.sort(key=lambda r: r.score, reverse=True)
        return results[:top_k]
    
    async def get_context(
        self,
        query: str,
        max_context: int = 4000,
        top_k: int = 3,
    ) -> str:
        """Get relevant context for a query"""
        results = await self.search(query, top_k=top_k)
        
        context_parts = []
        total_length = 0
        
        for result in results:
            if total_length + len(result.content) > max_context:
                break
            context_parts.append(result.content)
            total_length += len(result.content)
        
        return "\n\n---\n\n".join(context_parts)
    
    def _cosine_similarity(self, a: List[float], b: List[float]) -> float:
        """Calculate cosine similarity"""
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x * x for x in a) ** 0.5
        norm_b = sum(x * x for x in b) ** 0.5
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot / (norm_a * norm_b)


class StyleGuideRAG:
    """
    Specialized RAG for style guides and reference documents.
    
    Automatically loads and maintains context from:
    - Style guides
    - Character bibles
    - Brand guidelines
    - Writing examples
    """
    
    def __init__(
        self,
        registry: ModelRegistry,
        documents_path: str = "./documents",
    ):
        self.loader = RAGLoader()
        self.retriever = RAGRetriever(self.loader, registry)
        self.memory = HighContextMemory()
        self.documents_path = documents_path
    
    async def ingest(
        self,
        path: str,
        doc_type: str = "style_guide",
    ):
        """
        Ingest a document into the RAG system.
        
        Args:
            path: Path to document or directory
            doc_type: Type of document (style_guide, character_bank, etc.)
        """
        path_obj = Path(path)
        
        if path_obj.is_file():
            doc = await self.loader.load_file(path)
        else:
            docs = await self.loader.load_directory(path)
            if docs:
                doc = docs[0]
            else:
                return
        
        await self.retriever.index_documents([doc])
        
        # Also add to memory for direct access
        if doc_type == "style_guide":
            self.memory.set_style_guide({
                "source": doc.path,
                "content": doc.content[:10000],  # First 10K chars
            })
        elif doc_type == "character_bank":
            # Parse character info from content
            pass
    
    async def get_style_context(
        self,
        query: str,
    ) -> str:
        """Get style-relevant context"""
        results = await self.retriever.search(query, top_k=3)
        
        context = []
        for r in results:
            context.append(f"[Style Reference - Score: {r.score:.2f}]\n{r.content}")
        
        return "\n\n".join(context)
